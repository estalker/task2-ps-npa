from __future__ import annotations

from pathlib import Path

from docx import Document


def _fix_mojibake(text: str) -> str:
    try:
        from ftfy import fix_text  # type: ignore

        return fix_text(text)
    except Exception:
        return text


def extract_text_from_docx(path: Path) -> str:
    # Prefer mammoth: often better for "real-world" docx exports.
    try:
        import mammoth  # type: ignore

        with path.open("rb") as f:
            result = mammoth.extract_raw_text(f)
        txt = _fix_mojibake((result.value or "").strip())
        if txt:
            return txt
    except Exception:
        pass

    doc = Document(str(path))

    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # basic table extraction (often used in profstandards)
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                t = " ".join((cell.text or "").split())
                if t:
                    cells.append(t)
            if cells:
                parts.append(" | ".join(cells))

    return _fix_mojibake("\n".join(parts).strip())

